import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as exc:
    raise SystemExit("Missing dependency 'python-docx'. Please install it: pip install python-docx") from exc


OUTPUT_NAME = "SAM_Project_Documentation.docx"


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False, italic: bool = False, alignment: str | None = None):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if alignment == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def add_bullets(document: Document, items: list[str]):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]):
    for item in items:
        document.add_paragraph(item, style="List Number")


def generate_doc(output_path: str):
    document = Document()

    # Cover
    add_heading(document, "Segment Anything Model (SAM) - Project Documentation", level=0)
    add_paragraph(document, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", italic=True)
    add_paragraph(document, "Audience: Project Manager, Software Engineer", italic=True)

    # Executive Summary (PM)
    add_heading(document, "1. Executive Summary", level=1)
    add_paragraph(document, (
        "This project integrates Meta AI's Segment Anything Model (SAM) to enable fast, flexible "
        "image segmentation. It includes a Python backend library and a TypeScript demo UI for "
        "interactive segmentation and ONNX inference."))
    add_bullets(document, [
        "Outcome: Generate high-quality object masks from images with minimal prompts",
        "Deliverables: Python package, demo web app, example notebooks, output masks",
        "Stakeholders: PM, ML/Software Engineers",
        "Status: Ready for local runs and demonstrations",
    ])

    # Scope and Objectives
    add_heading(document, "2. Scope & Objectives", level=1)
    add_bullets(document, [
        "Integrate SAM core components for promptable segmentation",
        "Provide automatic mask generation utility",
        "Export model to ONNX and demo browser inference",
        "Offer notebooks for common usage patterns",
    ])

    # Architecture Overview
    add_heading(document, "3. Architecture Overview", level=1)
    add_paragraph(document, "High-level components:")
    add_numbered(document, [
        "Python package (`segment_anything/`): model, encoders, decoders, utilities",
        "Scripts (`scripts/`): automatic mask generator, ONNX export",
        "Notebooks (`notebooks/`): runnable examples",
        "Web demo (`segment-anything/demo/`): React + TypeScript UI",
        "Artifacts (`output_dir/`): generated masks and overlays",
    ])

    # Key Directories
    add_heading(document, "4. Key Directories & Files", level=1)
    add_bullets(document, [
        "segment_anything/modeling/sam.py: SAM model assembly",
        "segment_anything/predictor.py: High-level predictor API",
        "segment_anything/automatic_mask_generator.py: Batch/auto mask generation",
        "segment_anything/utils/onnx.py: ONNX export helpers",
        "scripts/amg.py: CLI for automatic mask generation",
        "scripts/export_onnx_model.py: Export to ONNX",
        "segment-anything/demo/src: Frontend demo application",
        "output_dir/mask/*.png: Generated masks",
        "output_dir/mask_overlay.png: Composite overlay",
    ])

    # Installation & Setup
    add_heading(document, "5. Setup & Installation", level=1)
    add_paragraph(document, "Prerequisites:")
    add_bullets(document, [
        "Python 3.9+ with pip",
        "Node.js 16+ (for web demo)",
        "GPU optional; CPU works for demos (slower)",
    ])
    add_paragraph(document, "Python package install:")
    document.add_paragraph("pip install -e .", style="Intense Quote")
    add_paragraph(document, "Optional extras: PyTorch, onnxruntime, onnxruntime-gpu depending on environment.")

    # Usage - Python
    add_heading(document, "6. Usage - Python", level=1)
    add_paragraph(document, "Predictor example (points, boxes, masks): see `notebooks/predictor_example.ipynb`.")
    add_paragraph(document, "Automatic mask generation:")
    document.add_paragraph(
        "python scripts/amg.py --input <image_or_folder> --output output_dir/",
        style="Intense Quote",
    )

    # Usage - ONNX & Web Demo
    add_heading(document, "7. ONNX & Web Demo", level=1)
    add_paragraph(document, "Export ONNX model:")
    document.add_paragraph(
        "python scripts/export_onnx_model.py --checkpoint <path_to_ckpt> --output sam.onnx",
        style="Intense Quote",
    )
    add_paragraph(document, "Run demo app:")
    document.add_paragraph(
        "cd segment-anything/demo && npm ci && npm run dev",
        style="Intense Quote",
    )

    # Data Flow
    add_heading(document, "8. Data Flow", level=1)
    add_bullets(document, [
        "Input image loaded and preprocessed (transforms)",
        "Prompt encoding (points/boxes/masks)",
        "Image encoding via vision transformer",
        "Mask decoding conditioned on prompts",
        "Post-processing and thresholding to produce binary masks",
    ])

    # Quality, Performance, and Risks
    add_heading(document, "9. Quality, Performance, and Risks", level=1)
    add_bullets(document, [
        "Deterministic inference given fixed prompts",
        "Performance depends on model size and hardware",
        "Memory use may be high for large images",
        "ONNX runtime on CPU may be slow; prefer GPU if available",
    ])

    # Roadmap
    add_heading(document, "10. Roadmap & Next Steps", level=1)
    add_bullets(document, [
        "Add CLI for promptable single-image segmentation",
        "Dockerize demo and backend for easier deployment",
        "Provide evaluation scripts and metrics",
        "Integrate continuous testing and linting in CI",
    ])

    # Appendix
    add_heading(document, "Appendix A: References", level=1)
    add_bullets(document, [
        "README.md for overall guidance",
        "Notebooks for hands-on examples",
        "Demo `README.md` for web UI",
    ])

    document.save(output_path)


if __name__ == "__main__":
    out_path = os.path.join(os.path.dirname(__file__), OUTPUT_NAME)
    generate_doc(out_path)
    print(f"Wrote documentation to: {out_path}")


